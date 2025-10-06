import re
import base64
import io
from typing import Dict, List, Optional, Any
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract

class CVParser:
    """Parser for extracting information from CV files"""
    
    def __init__(self):
        self.email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        self.phone_pattern = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{2,4}[-.\s]?\d{2,4}[-.\s]?\d{0,4}'
        self.linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
        
        # Common keywords for extraction
        self.education_keywords = ['diplôme', 'formation', 'éducation', 'université', 'école', 'master', 'licence', 'bac']
        self.experience_keywords = ['expérience', 'emploi', 'poste', 'travail', 'entreprise', 'société']
        self.skills_keywords = ['compétences', 'compétence', 'skills', 'technologies', 'langages']
        self.languages_keywords = ['langues', 'langue', 'languages']
        
    def parse_file(self, file_content: bytes, filename: str, file_type: str) -> Dict[str, Any]:
        """Main parser that routes to appropriate parser based on file type"""
        try:
            if file_type == 'application/pdf':
                return self.parse_pdf(file_content)
            elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                return self.parse_docx(file_content)
            elif file_type.startswith('text/'):
                return self.parse_text(file_content)
            elif file_type.startswith('image/'):
                return self.parse_image(file_content, file_type)
            else:
                return {"error": "Format de fichier non supporté"}
        except Exception as e:
            return {"error": f"Erreur lors de l'analyse du CV: {str(e)}"}
    
    def parse_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text and data from PDF files"""
        try:
            text = ""
            images = []
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    
                    # Try to extract images
                    if hasattr(page, 'images'):
                        for img in page.images:
                            images.append(img)
            
            # Extract structured data from text
            extracted_data = self.extract_data_from_text(text)
            
            # Try to extract photo
            if images and len(images) > 0:
                try:
                    # Get first image as potential candidate photo
                    extracted_data['has_photo'] = True
                except:
                    pass
            
            extracted_data['cv_original'] = base64.b64encode(file_content).decode('utf-8')
            
            return extracted_data
            
        except Exception as e:
            return {"error": f"Erreur lors de l'analyse du PDF: {str(e)}"}
    
    def parse_docx(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text and data from DOCX files"""
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract data from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            extracted_data = self.extract_data_from_text(text)
            extracted_data['cv_original'] = base64.b64encode(file_content).decode('utf-8')
            
            # Try to extract images from document
            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        extracted_data['has_photo'] = True
                        break
            except:
                pass
            
            return extracted_data
            
        except Exception as e:
            return {"error": f"Erreur lors de l'analyse du DOCX: {str(e)}"}
    
    def parse_text(self, file_content: bytes) -> Dict[str, Any]:
        """Extract data from plain text files"""
        try:
            text = file_content.decode('utf-8', errors='ignore')
            extracted_data = self.extract_data_from_text(text)
            extracted_data['cv_original'] = base64.b64encode(file_content).decode('utf-8')
            return extracted_data
        except Exception as e:
            return {"error": f"Erreur lors de l'analyse du texte: {str(e)}"}
    
    def parse_image(self, file_content: bytes, file_type: str) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        try:
            image = Image.open(io.BytesIO(file_content))
            
            # Store image as photo
            photo_base64 = base64.b64encode(file_content).decode('utf-8')
            
            # Try OCR to extract text
            try:
                text = pytesseract.image_to_string(image, lang='fra+eng')
            except:
                text = ""
            
            extracted_data = self.extract_data_from_text(text) if text else {}
            extracted_data['photo_base64'] = photo_base64
            extracted_data['has_photo'] = True
            extracted_data['cv_original'] = photo_base64
            
            return extracted_data
            
        except Exception as e:
            return {"error": f"Erreur lors de l'analyse de l'image: {str(e)}"}
    
    def extract_data_from_text(self, text: str) -> Dict[str, Any]:
        """Extract structured data from raw text"""
        data = {
            'auto_filled_fields': []
        }
        
        # Extract email
        email_matches = re.findall(self.email_pattern, text, re.IGNORECASE)
        if email_matches:
            data['email'] = email_matches[0]
            data['auto_filled_fields'].append('email')
        
        # Extract phone
        phone_matches = re.findall(self.phone_pattern, text)
        if phone_matches:
            # Clean phone number
            phone = re.sub(r'[^\d+]', '', phone_matches[0])
            data['telephone'] = phone_matches[0]
            data['auto_filled_fields'].append('telephone')
        
        # Extract LinkedIn
        linkedin_matches = re.findall(self.linkedin_pattern, text, re.IGNORECASE)
        if linkedin_matches:
            data['profil_linkedin'] = linkedin_matches[0]
            data['auto_filled_fields'].append('profil_linkedin')
        
        # Extract name (first line often contains name)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if lines:
            first_line = lines[0]
            # Assume first line with 2-3 words is the name
            words = first_line.split()
            if 2 <= len(words) <= 4 and not any(char.isdigit() for char in first_line):
                if len(words) >= 2:
                    data['prenom'] = words[0]
                    data['nom'] = ' '.join(words[1:])
                    data['auto_filled_fields'].extend(['prenom', 'nom'])
        
        # Extract competences (skills)
        competences = self.extract_skills(text)
        if competences:
            data['competences'] = competences
            data['auto_filled_fields'].append('competences')
        
        # Extract langues (languages)
        langues = self.extract_languages(text)
        if langues:
            data['langues'] = langues
            data['auto_filled_fields'].append('langues')
        
        # Extract experiences
        experiences = self.extract_experiences(text)
        if experiences:
            data['experiences'] = experiences
            data['auto_filled_fields'].append('experiences')
        
        # Extract formations
        formations = self.extract_formations(text)
        if formations:
            data['formations'] = formations
            data['auto_filled_fields'].append('formations')
        
        # Extract current position
        poste = self.extract_current_position(text)
        if poste:
            data['poste_actuel'] = poste
            data['auto_filled_fields'].append('poste_actuel')
        
        return data
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from text"""
        skills = []
        text_lower = text.lower()
        
        # Common technical skills
        common_skills = [
            'python', 'java', 'javascript', 'react', 'angular', 'vue', 'node.js', 'nodejs',
            'php', 'laravel', 'django', 'flask', 'fastapi', 'html', 'css', 'sql', 'mysql',
            'postgresql', 'mongodb', 'docker', 'kubernetes', 'aws', 'azure', 'git', 'github',
            'tailwind', 'bootstrap', 'typescript', 'c++', 'c#', 'ruby', 'go', 'rust',
            'machine learning', 'ai', 'data science', 'agile', 'scrum'
        ]
        
        for skill in common_skills:
            if skill in text_lower:
                skills.append(skill.title())
        
        return list(set(skills))[:15]  # Limit to 15 skills
    
    def extract_languages(self, text: str) -> List[Dict[str, str]]:
        """Extract languages from text"""
        languages = []
        text_lower = text.lower()
        
        language_patterns = [
            (r'français|french', 'Français'),
            (r'anglais|english', 'Anglais'),
            (r'arabe|arabic', 'Arabe'),
            (r'espagnol|spanish', 'Espagnol'),
            (r'allemand|german', 'Allemand'),
            (r'italien|italian', 'Italien')
        ]
        
        for pattern, lang_name in language_patterns:
            if re.search(pattern, text_lower):
                # Try to find level
                niveau = "Non spécifié"
                if re.search(rf'{pattern}.*?(courant|fluent|bilingue|natif)', text_lower):
                    niveau = "Courant"
                elif re.search(rf'{pattern}.*?(intermédiaire|intermediate)', text_lower):
                    niveau = "Intermédiaire"
                elif re.search(rf'{pattern}.*?(débutant|basic|basique)', text_lower):
                    niveau = "Débutant"
                
                languages.append({
                    'langue': lang_name,
                    'niveau': niveau
                })
        
        return languages
    
    def extract_experiences(self, text: str) -> List[Dict[str, str]]:
        """Extract professional experiences from text"""
        experiences = []
        
        # This is a simplified extraction - in production, use more sophisticated NLP
        lines = text.split('\n')
        in_experience_section = False
        current_exp = {}
        
        for line in lines:
            line_lower = line.lower()
            
            # Detect experience section
            if any(keyword in line_lower for keyword in self.experience_keywords):
                in_experience_section = True
                continue
            
            # Exit experience section
            if in_experience_section and any(keyword in line_lower for keyword in self.education_keywords + self.skills_keywords):
                in_experience_section = False
                if current_exp:
                    experiences.append(current_exp)
                    current_exp = {}
                continue
            
            if in_experience_section and line.strip():
                # Try to detect dates
                date_pattern = r'\d{4}|\d{2}/\d{4}'
                if re.search(date_pattern, line):
                    if current_exp:
                        experiences.append(current_exp)
                    current_exp = {
                        'poste': line.strip(),
                        'entreprise': '',
                        'date_debut': '',
                        'date_fin': '',
                        'description': ''
                    }
        
        if current_exp:
            experiences.append(current_exp)
        
        return experiences[:5]  # Limit to 5 experiences
    
    def extract_formations(self, text: str) -> List[Dict[str, str]]:
        """Extract education/training from text"""
        formations = []
        
        lines = text.split('\n')
        in_education_section = False
        current_formation = {}
        
        for line in lines:
            line_lower = line.lower()
            
            # Detect education section
            if any(keyword in line_lower for keyword in self.education_keywords):
                in_education_section = True
                continue
            
            # Exit education section
            if in_education_section and any(keyword in line_lower for keyword in self.experience_keywords + self.skills_keywords):
                in_education_section = False
                if current_formation:
                    formations.append(current_formation)
                    current_formation = {}
                continue
            
            if in_education_section and line.strip():
                # Try to detect dates
                date_pattern = r'\d{4}|\d{2}/\d{4}'
                if re.search(date_pattern, line):
                    if current_formation:
                        formations.append(current_formation)
                    current_formation = {
                        'diplome': line.strip(),
                        'etablissement': '',
                        'date_debut': '',
                        'date_fin': '',
                        'description': ''
                    }
        
        if current_formation:
            formations.append(current_formation)
        
        return formations[:5]  # Limit to 5 formations
    
    def extract_current_position(self, text: str) -> Optional[str]:
        """Extract current position from text"""
        # Look for common patterns
        patterns = [
            r'poste actuel[:\s]+([^\n]+)',
            r'current position[:\s]+([^\n]+)',
            r'titre[:\s]+([^\n]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
